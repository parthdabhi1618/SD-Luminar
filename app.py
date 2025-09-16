from flask import Flask, render_template, request, jsonify, Response, send_from_directory
from flask_wtf.csrf import CSRFProtect
import fitz
import os
import tempfile
import re
import subprocess
import html
import uuid
from werkzeug.utils import secure_filename
import threading

# ReportLab Imports
from reportlab.platypus import BaseDocTemplate, Paragraph, Spacer, Frame, PageTemplate, Preformatted, Image
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from io import BytesIO
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt, Inches

# Pygments, Matplotlib, yt-dlp, and Scraping libraries
import matplotlib.pyplot as plt
from youtube_downloader import YouTubeDownloader
import requests
import time
try:
    import yt_dlp
except Exception:
    yt_dlp = None

app = Flask(__name__)
temp_dir = tempfile.mkdtemp()
app.config['UPLOAD_FOLDER'] = temp_dir
app.config['SECRET_KEY'] = 'luminar-secret-key-2025'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
plt.switch_backend('agg')
csrf = CSRFProtect(app)

# --- UTILITY & CORE LOGIC FUNCTIONS ---
dark_bg = colors.HexColor("#0A0A0A")
white_text = colors.HexColor("#FFFFFF")
green_accent = colors.HexColor("#00FF41")
cyan_accent = colors.HexColor("#00D4FF")

def to_roman(n):
    if not isinstance(n, int) or n <= 0: return str(n)
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syb = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman_num = ''
    i = 0
    while n > 0:
        for _ in range(n // val[i]):
            roman_num += syb[i]
            n -= val[i]
        i += 1
    return roman_num

def to_alpha(n, uppercase=True):
    if not isinstance(n, int) or n <= 0: return str(n)
    result = ""
    start = 65 if uppercase else 97
    while n > 0:
        n, remainder = divmod(n - 1, 26)
        result = chr(start + remainder) + result
    return result

def get_doc_stats(filepath):
    try:
        doc = fitz.open(filepath)
        page_count = len(doc)
        text = "".join(page.get_text() for page in doc)
        char_count = len(text)
        word_count = len(text.split())
        doc.close()
        return {'pages': page_count, 'words': word_count, 'characters': char_count}
    except Exception as e:
        print(f"Could not get stats for {filepath}: {e}")
        return {'pages': 0, 'words': 0, 'characters': 0}


# Background conversion tracking: map server_filename -> {'status': 'pending'|'done'|'failed', 'pdf_path': str}
conversion_status = {}

@app.route('/check_conversion_status/<filename>')
def check_conversion_status(filename):
    # Try direct lookup first (filename might be the original server-side name)
    status = conversion_status.get(filename)

    # If not found, filename might be a pdf basename (e.g. converted file). Try reverse lookup.
    if status is None and filename and filename.lower().endswith('.pdf'):
        for k, v in conversion_status.items():
            try:
                if v and v.get('pdf_path') and os.path.basename(v.get('pdf_path')) == filename:
                    status = dict(v)
                    # expose the original server key too for convenience
                    status['serverFilename'] = k
                    break
            except Exception:
                continue

    # Fallback to unknown if still nothing
    if status is None:
        status = {'status': 'unknown'}

    # Ensure frontend-friendly fields
    if isinstance(status, dict) and status.get('pdf_path'):
        try:
            status = dict(status)
            status['pdf_basename'] = os.path.basename(status.get('pdf_path'))
        except Exception:
            pass
    return jsonify(status)

def _convert_ipynb_to_pdf_async(src_path, timeout=120):
    """Run nbconvert to webpdf in a background thread and update conversion_status.

    This implementation uses Popen so we can provide a simple time-based progress
    estimate while nbconvert runs. nbconvert does not expose a machine-readable
    progress API here, so we keep a lightweight heuristic: while the process is
    running we increment progress periodically and then set status to done on
    success or failed on error.
    """
    server_key = os.path.basename(src_path)
    pdf_output_path = os.path.splitext(src_path)[0] + '.pdf'
    try:
        # Mark started
        conversion_status[server_key] = {'status': 'pending', 'pdf_path': None, 'progress': 0}

        # Start nbconvert as a subprocess and poll
        proc = subprocess.Popen(['jupyter', 'nbconvert', '--to', 'webpdf', '--allow-chromium-download', src_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)

        # Simple progress estimator: ramp from 5 -> 85 while running
        progress = 5
        conversion_status[server_key]['progress'] = progress

        # Poll process and update progress
        while proc.poll() is None:
            time.sleep(1)
            # increase progress slowly but cap below 95
            progress = min(85, progress + 5)
            conversion_status[server_key]['progress'] = progress

        # Process finished, check return code
        rc = proc.returncode
        if rc == 0 and os.path.exists(pdf_output_path):
            conversion_status[server_key] = {'status': 'done', 'pdf_path': pdf_output_path, 'progress': 100}
            try:
                conversion_status[server_key]['pdf_basename'] = os.path.basename(pdf_output_path)
            except Exception:
                pass
            print(f"Background conversion completed: {pdf_output_path}")
        else:
            # grab stderr for diagnostics
            try:
                _, err = proc.communicate(timeout=1)
                print(err.decode('utf-8', errors='ignore'))
            except Exception:
                pass
            conversion_status[server_key] = {'status': 'failed', 'pdf_path': None, 'progress': 0}
            print(f"Background conversion failed for {src_path} (rc={rc})")
    except Exception as e:
        conversion_status[server_key] = {'status': 'failed', 'pdf_path': None, 'progress': 0}
        print(f"Background nbconvert failed for {src_path}: {e}")


def get_pdf_for_serverfile(server_filename, input_filepath, wait_seconds=5, sync_timeout=60):
    """Return a PDF path for the given server_filename.

    Logic:
      - If the uploaded file is already a PDF, return it.
      - If conversion_status shows a done PDF, return that path.
      - If conversion_status shows pending, wait up to wait_seconds for it to finish.
      - If still not available, try a synchronous nbconvert attempt (sync_timeout).
      - Return None if no PDF could be obtained.
    """
    # If caller already passed a PDF basename or filename, prefer that file if it exists
    if server_filename and server_filename.lower().endswith('.pdf'):
        candidate = os.path.join(app.config['UPLOAD_FOLDER'], os.path.basename(server_filename))
        if os.path.exists(candidate):
            return candidate

    # Already a PDF (full path passed)
    if server_filename and os.path.exists(input_filepath) and server_filename.lower().endswith('.pdf'):
        return input_filepath

    # Look up conversion status by the exact server_filename key
    status = conversion_status.get(server_filename)

    # If not found, it's possible the caller passed the converted PDF basename. Try reverse lookup.
    if status is None and server_filename and server_filename.lower().endswith('.pdf'):
        for k, v in conversion_status.items():
            try:
                if v and v.get('pdf_path') and os.path.basename(v.get('pdf_path')) == os.path.basename(server_filename):
                    status = v
                    break
            except Exception:
                continue
    if status and status.get('status') == 'done' and status.get('pdf_path') and os.path.exists(status.get('pdf_path')):
        return status.get('pdf_path')

    # If pending, wait briefly
    if status and status.get('status') == 'pending':
        waited = 0.0
        while waited < wait_seconds:
            time.sleep(0.5)
            waited += 0.5
            status = conversion_status.get(server_filename)
            if status and status.get('status') == 'done' and status.get('pdf_path') and os.path.exists(status.get('pdf_path')):
                return status.get('pdf_path')

    # Try synchronous conversion as a fallback
    pdf_output_path = os.path.splitext(input_filepath)[0] + '.pdf'
    try:
        # Mark pending so frontend knows conversion started
        conversion_status[server_filename] = {'status': 'pending', 'pdf_path': None, 'progress': 0}
        subprocess.run(['jupyter', 'nbconvert', '--to', 'webpdf', '--allow-chromium-download', input_filepath], check=True, timeout=sync_timeout)
        if os.path.exists(pdf_output_path):
            conversion_status[server_filename] = {'status': 'done', 'pdf_path': pdf_output_path, 'progress': 100}
            return pdf_output_path
    except Exception as e:
        print(f"Synchronous nbconvert fallback failed for {input_filepath}: {e}")
        conversion_status[server_filename] = {'status': 'failed', 'pdf_path': None, 'progress': 0}

    return None

def extract_highlights(pdf_path):
    from highlight_extractor import HighlightExtractor
    doc = fitz.open(pdf_path)
    extractor = HighlightExtractor(doc)
    highlights = extractor.extract_highlights()

    # Process pages in order and collect highlights with page context
    for page_num, page in enumerate(doc):
        page_highlights = []
        for annot in page.annots():
            if annot.type[1] == "Highlight":
                rect = annot.rect
                words = page.get_text("words", clip=rect)
                if not words: continue

                # Sort words by their position (left to right, top to bottom)
                words_sorted = sorted(words, key=lambda w: (round(w[3], 1), w[0]))  # Sort by y_pos, then x_pos

                # Group words into lines based on y-position
                lines = {}
                for word in words_sorted:
                    y_pos = round(word[3], 1)
                    if y_pos not in lines:
                        lines[y_pos] = []
                    lines[y_pos].append(word[4])

                # Process lines in reading order (top to bottom)
                for y_pos in sorted(lines.keys()):
                    line_text = ' '.join(lines[y_pos]).strip()
                    if not line_text: continue

                    # Add page context to maintain document order
                    highlight_data = {
                        'text': line_text,
                        'page': page_num + 1,
                        'y_pos': y_pos
                    }
                    page_highlights.append(highlight_data)

        # Sort page highlights by vertical position and add to main list
        page_highlights.sort(key=lambda h: h['y_pos'])
        for highlight in page_highlights:
            line_text = highlight['text']

            # Improved categorization logic
            category = categorize_highlight(line_text)
            highlights.append((category, line_text))

    doc.close()
    return highlights

def categorize_highlight(text):
    """Improved categorization logic for highlights"""
    text = text.strip()

    # Check for headings first (most specific)
    if re.match(r'^(Chapter|Section|Topic|Lesson|Module|Unit|Part)\s+\d+', text, re.IGNORECASE):
        return 'heading'
    if re.match(r'^(AIM:|Objective:|Goal:|Learning Objective:|Key Concept:)', text, re.IGNORECASE):
        return 'heading'
    if len(text) < 80 and text.istitle() and not any(char in text for char in '.,!?;:'):
        return 'heading'

    # Check for code (programming languages)
    code_keywords = [
        # Python
        r'\b(def|class|import|from|if|elif|else|for|while|try|except|with|as|lambda|return|yield)\b',
        # Java/JavaScript
        r'\b(public|private|protected|static|void|int|String|function|var|let|const|class|interface)\b',
        # C/C++/C#
        r'\b(int|char|float|double|void|struct|class|public|private|protected|static)\b',
        # SQL
        r'\b(SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|FROM|WHERE|JOIN|GROUP BY|ORDER BY)\b',
        # General programming
        r'\b(print|console\.log|System\.out\.println)\b'
    ]

    code_indicators = [
        r'[{}();=<>]',  # Common programming symbols
        r'\[.*\]',      # Array/list access
        r'\(.*\)\s*{',  # Function definitions
        r'import\s+.*', # Import statements
        r'#include',    # C/C++ includes
    ]

    for keyword in code_keywords:
        if re.search(keyword, text, re.IGNORECASE):
            return 'code'

    for indicator in code_indicators:
        if re.search(indicator, text):
            return 'code'

    # Check for mathematical expressions
    math_patterns = [
        r'[+\-×÷=≠≈≤≥∞∑∫√∛∜∂∇∆∅∈∉⊂⊃∪∩∧∨¬⇒⇔∀∃∄]',  # Mathematical symbols
        r'\b(sin|cos|tan|log|ln|exp|sqrt|pi|e|alpha|beta|gamma|delta)\b',  # Math functions/constants
        r'\d+\s*[+\-×÷=]\s*\d+',  # Simple arithmetic
        r'\b\d+\^\d+\b',  # Exponents
        r'\b\d+/\d+\b',   # Fractions
        r'\(\d+\)',       # Parenthesized numbers
    ]

    for pattern in math_patterns:
        if re.search(pattern, text):
            return 'math'

    # Check for lists and bullet points
    if re.match(r'^[-•*]\s', text):
        return 'list_item'

    # Check for numbered lists
    if re.match(r'^\d+[\.)]\s', text):
        return 'list_item'

    # Check for questions
    if text.endswith('?') or text.startswith(('What', 'How', 'Why', 'When', 'Where', 'Who')):
        return 'question'

    # Check for important terms or definitions
    if ':' in text and len(text.split(':')[0].strip()) < 30:
        return 'definition'

    # Default to regular text
    return 'point'

def create_modern_pdf(highlights, output_path):
    left_margin, right_margin, top_margin, bottom_margin = (0.75*inch,) * 4

    class ModernDocTemplate(BaseDocTemplate):
        def __init__(self, filename, **kw):
            super().__init__(filename, **kw)
            frame = Frame(left_margin, bottom_margin, self.width, self.height, id='content_frame')
            template = PageTemplate(id='main', frames=[frame], onPage=self.page_template)
            self.addPageTemplates([template])

        def page_template(self, canvas, doc):
            canvas.saveState()
            # Clean white background
            canvas.setFillColor(colors.white)
            canvas.rect(0, 0, self.pagesize[0], self.pagesize[1], fill=1)

            # Modern header with Hephaestus branding
            canvas.setFillColor(colors.HexColor("#4a90e2"))  # Professional blue
            canvas.setFont('Helvetica-Bold', 12)
            canvas.drawCentredString(self.pagesize[0]/2, self.pagesize[1] - 0.5*inch, "Hephaestus")

            # Clean page number
            canvas.setFillColor(colors.gray)
            canvas.setFont('Helvetica', 9)
            canvas.drawCentredString(self.pagesize[0]/2, 0.5*inch, f"Page {doc.page}")
            canvas.restoreState()

    doc = ModernDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()

    # Modern, clean styles
    styles.add(ParagraphStyle(
        name='ModernHeading',
        fontName='Helvetica-Bold',
        fontSize=16,
        textColor=colors.HexColor("#2c3e50"),  # Dark blue-gray
        alignment=TA_CENTER,
        spaceAfter=15,
        leading=20
    ))

    styles.add(ParagraphStyle(
        name='ModernBody',
        fontName='Helvetica',
        fontSize=12,
        textColor=colors.black,
        spaceAfter=10,
        leading=16,
        firstLineIndent=0
    ))

    styles.add(ParagraphStyle(
        name='ModernCode',
        fontName='Courier',
        fontSize=11,
        textColor=colors.HexColor("#2c3e50"),
        backColor=colors.HexColor("#f8f9fa"),  # Light gray background
        borderColor=colors.HexColor("#dee2e6"),
        borderWidth=1,
        borderPadding=8,
        leading=14,
        spaceAfter=12
    ))

    story = []

    # Add title
    title = Paragraph("Study Notes", styles['ModernHeading'])
    story.append(title)
    story.append(Spacer(1, 20))

    # Process highlights without table of contents
    for item_type, text in highlights:
        if item_type == 'heading':
            p = Paragraph(text, styles['ModernHeading'])
            story.append(p)
        elif item_type == 'code':
            p = Preformatted(text, styles['ModernCode'])
            story.append(p)
        elif item_type == 'math':
            try:
                fig = plt.figure(figsize=(6, 1), facecolor='white')
                fig.text(0.5, 0.5, f'${text}$', ha='center', va='center', fontsize=20, color='black')
                img_path = os.path.join(temp_dir, f'math_{hash(text)}.png')
                plt.savefig(img_path, transparent=True, bbox_inches='tight', pad_inches=0.1)
                plt.close(fig)
                story.append(Image(img_path, width=4*inch, height=0.5*inch))
            except Exception:
                story.append(Paragraph(text, styles['ModernBody']))
        else:
            # Clean bullet points
            p_text = f'<bullet color="#4a90e2">•</bullet> {html.escape(text)}'
            story.append(Paragraph(p_text, styles['ModernBody']))

        story.append(Spacer(1, 8))

    doc.build(story)

def create_docx_from_highlights(highlights, output_path):
    doc = Document()
    styles = doc.styles
    if 'Normal' in styles:
        styles['Normal'].font.name = 'Calibri'
        styles['Normal'].font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    run = title.add_run('Study Notes')
    run.bold = True
    run.font.size = Pt(18)

    for item_type, text in highlights:
        if item_type == 'heading':
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(14)
        elif item_type == 'code':
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        else:
            p = doc.add_paragraph('• ' + text)
            p_format = p.paragraph_format
            p_format.space_after = Pt(6)
    doc.save(output_path)

def create_docx_from_pdf(pdf_path: str, docx_output_path: str) -> None:
    try:
        pdf = fitz.open(pdf_path)
        doc = Document()
        for page in pdf:
            pix = page.get_pixmap(dpi=144)
            img_path = os.path.join(temp_dir, f"page_{uuid.uuid4()}.png")
            pix.save(img_path)
            doc.add_picture(img_path, width=Inches(6.5))
        doc.save(docx_output_path)
    except Exception as e:
        print(f"Failed to create DOCX from PDF: {e}")

def add_header_footer_to_pdf(input_pdf_path, output_filepath, headers, footers, start_page_num, page_num_placement, page_num_format, overlap_resolution, margin_size, chapter_num, page_num_enabled, hf_enabled):
    input_doc = fitz.open(input_pdf_path)
    output_doc = fitz.open()
    header_y_pos = letter[1] - 0.5*inch
    footer_y_pos = 0.5*inch
    hf_x_margin = 0.5*inch
    # Default to narrow margins (0.1 inch) without user selection
    content_margin = 0.1*inch
    total_pages = len(input_doc)
    last_page_num = start_page_num + total_pages - 1
    page_num_area, page_num_pos = page_num_placement.split('-')
    for i, page in enumerate(input_doc):
        temp_headers = headers.copy()
        temp_footers = footers.copy()
        new_page = output_doc.new_page(width=page.rect.width, height=page.rect.height)
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=(page.rect.width, page.rect.height))
        can.setFont('Helvetica', 9)
        current_page_num = start_page_num + i
        page_num_str = ""
        if page_num_enabled:
            format_map = { 'roman_lower': to_roman(current_page_num).lower(), 'roman_upper': to_roman(current_page_num), 'alpha_lower': to_alpha(current_page_num, False), 'alpha_upper': to_alpha(current_page_num), 'dash_x_dash': f"- {current_page_num} -", 'page_x': f"Page {current_page_num}", 'page_x_of_n': f"Page {current_page_num} of {last_page_num}", 'book_style': f"{chapter_num}-{current_page_num}", }
            page_num_str = format_map.get(page_num_format, str(current_page_num))
            target_dict = temp_headers if page_num_area == 'header' else temp_footers
            pos_key = page_num_pos
            if target_dict.get(pos_key):
                if overlap_resolution == 'before': target_dict[pos_key] = f"{page_num_str} {target_dict[pos_key]}"
                else: target_dict[pos_key] = f"{target_dict[pos_key]} {page_num_str}"
                page_num_str = ""
        if hf_enabled:
            can.drawString(hf_x_margin, header_y_pos, temp_headers.get('left', ''))
            can.drawCentredString(page.rect.width / 2, header_y_pos, temp_headers.get('center', ''))
            can.drawRightString(page.rect.width - hf_x_margin, header_y_pos, temp_headers.get('right', ''))
            can.drawString(hf_x_margin, footer_y_pos, temp_footers.get('left', ''))
            can.drawCentredString(page.rect.width / 2, footer_y_pos, temp_footers.get('center', ''))
            can.drawRightString(page.rect.width - hf_x_margin, footer_y_pos, temp_footers.get('right', ''))
        if page_num_str:
            y_pos = header_y_pos if page_num_area == 'header' else footer_y_pos
            if page_num_pos == 'left': can.drawString(hf_x_margin, y_pos, page_num_str)
            elif page_num_pos == 'right': can.drawRightString(page.rect.width - hf_x_margin, y_pos, page_num_str)
            else: can.drawCentredString(page.rect.width / 2, y_pos, page_num_str)
        can.save()
        packet.seek(0)
        overlay_doc = fitz.open("pdf", packet.read())
        new_page.show_pdf_page(new_page.rect, overlay_doc, 0)
        content_rect = fitz.Rect(content_margin, content_margin, page.rect.width - content_margin, page.rect.height - content_margin)
        new_page.show_pdf_page(content_rect, input_doc, i)
    output_doc.save(output_filepath)
    output_doc.close()
    input_doc.close()

# --- FLASK ROUTES ---
@app.route('/')
def home():
    return send_from_directory('templates', 'index.html')

@app.route('/temp/<path:filename>')
def serve_temp_file(filename):
    force_download = request.args.get('download') is not None or request.args.get('filename') is not None
    download_name = request.args.get('filename')
    try:
        return send_from_directory(temp_dir, filename, as_attachment=force_download, download_name=download_name)
    except TypeError:
        return send_from_directory(temp_dir, filename, as_attachment=force_download)

@csrf.exempt
@app.route('/download_youtube', methods=['POST', 'GET'])
def download_youtube():
    try:
        if request.method == 'POST':
            url = request.form.get('url')
            if not url:
                return jsonify({'error': 'No URL provided'}), 400
                
            # Initialize downloader with API key if available
            from smart_youtube_downloader import SmartYouTubeDownloader
            api_key = os.getenv('YOUTUBE_API_KEY')
            downloader = SmartYouTubeDownloader(api_key=api_key)
            
            # For initial request, just get preview info
            preview_info = downloader.get_preview_info(url)
            preview_info['download_id'] = str(uuid.uuid4())  # Add unique ID for tracking download
            return jsonify(preview_info)
            
        elif request.method == 'GET':
            # Handle actual download request
            url = request.args.get('url')
            download_id = request.args.get('download_id')
            if not url or not download_id:
                return jsonify({'error': 'Missing parameters'}), 400
            # NOTE: Actual streaming/download implementation goes here.
            # For now return a not-implemented response to close the try block safely.
            return jsonify({'error': 'Download-on-demand not implemented yet'}), 501
    except Exception as e:
        return jsonify({'error': f'Failed to process download request: {str(e)}'}), 500


@csrf.exempt
@app.route('/upload_and_analyze', methods=['POST'])
def upload_and_analyze():
    """Accept a file upload. If it's a PDF return immediate stats. If it's an
    IPYNB, save file, mark conversion pending and start background conversion.
    The frontend expects at least: serverFilename and initialStats/pageCount.
    """
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    f = request.files['file']
    if f.filename == '':
        return jsonify({'error': 'Empty filename'}), 400

    original_filename = secure_filename(f.filename)

    # Read uploaded bytes to compute a content-hash to deduplicate identical uploads
    file_bytes = f.read()
    import hashlib
    content_hash = hashlib.sha1(file_bytes).hexdigest()[:12]
    filename = f"{content_hash}_{original_filename}"
    server_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    # If file already exists with same content-hash, reuse it; otherwise write it
    if not os.path.exists(server_path):
        with open(server_path, 'wb') as outf:
            outf.write(file_bytes)
    else:
        # already present; do nothing (avoid creating duplicate files)
        pass

    # If PDF, return page stats right away
    if filename.lower().endswith('.pdf'):
        stats = get_doc_stats(server_path)
        return jsonify({'serverFilename': filename, 'initialStats': stats, 'pageCount': stats.get('pages', 1)})

    # If ipynb, mark pending and spawn background conversion
    if filename.lower().endswith('.ipynb'):
        # mark pending
        conversion_status[filename] = {'status': 'pending', 'pdf_path': None, 'progress': 0}
        # start thread
        t = threading.Thread(target=_convert_ipynb_to_pdf_async, args=(server_path, 300), daemon=True)
        t.start()
        # frontend will poll check_conversion_status
        return jsonify({'serverFilename': filename, 'initialStats': {'pages': 0}, 'pageCount': 0})

    # Unknown types: accept but return generic response
    return jsonify({'serverFilename': filename, 'initialStats': {}, 'pageCount': 0})


@app.route('/get_page_count')
def get_page_count():
    filename = request.args.get('filename')
    if not filename:
        return jsonify({'pageCount': 0}), 400
    try:
        # Try absolute path first
        path = filename if os.path.isabs(filename) else os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(path):
            # Try basename in temp dir
            path = os.path.join(app.config['UPLOAD_FOLDER'], os.path.basename(filename))
        if not os.path.exists(path):
            return jsonify({'pageCount': 0}), 404
        stats = get_doc_stats(path)
        return jsonify({'pageCount': stats.get('pages', 0)})
    except Exception as e:
        return jsonify({'pageCount': 0, 'error': str(e)}), 500

@csrf.exempt
@app.route('/video_download', methods=['POST'])
def video_download():
    url = request.json.get('url')
    if not url: return jsonify({'error': 'URL is required'}), 400

    # Enhanced yt-dlp options for better compatibility
    ydl_opts = {
        'quiet': True,
        'no_warnings': True,
        'extract_flat': False,
        'format_sort': ['res:1080', 'res:720', 'res:480', 'ext:mp4:m4a'],
        'prefer_free_formats': True,
    }

    # Make yt-dlp more tolerant for certain platforms (X/Twitter, geo-restricted) and reduce interactive prompts
    ydl_opts.update({
        'noplaylist': True,
        'geo_bypass': True,
        'nocheckcertificate': True,
        'ignoreerrors': True,
        'skip_unavailable_fragments': True,
    })

    try:
        if yt_dlp is None:
            return jsonify({'error': 'yt_dlp module not installed on server'}), 501
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            media_formats = []
            best_audio = None

            if 'formats' in info:
                for f in info.get('formats', []):
                    vcodec = f.get('vcodec')
                    acodec = f.get('acodec')
                    height = f.get('height', 0) or 0
                    ext = f.get('ext', 'mp4')
                    format_note = f.get('format_note', '')
                    quality = format_note if format_note else (f"{height}p" if height else ext.upper())
                    filesize = f.get('filesize_approx') or f.get('filesize') or 0
                    size_mb = round(filesize / (1024 * 1024), 1) if filesize else None

                    is_video_only = (vcodec and vcodec != 'none') and (not acodec or acodec == 'none')
                    is_audio_only = (acodec and acodec != 'none') and (not vcodec or vcodec == 'none')
                    is_progressive = (vcodec and vcodec != 'none') and (acodec and acodec != 'none')

                    entry = {
                        'quality': quality,
                        'url': f.get('url'),
                        'size': size_mb,
                        'format': ext,
                        'height': height,
                        'video_only': is_video_only,
                        'audio_only': is_audio_only,
                        'progressive': is_progressive,
                        'abr': f.get('abr'),
                        'vcodec': vcodec,
                        'acodec': acodec
                    }

                    if is_audio_only and (best_audio is None or (f.get('abr') or 0) > (best_audio.get('abr') or 0)):
                        best_audio = entry

                    if is_progressive or is_video_only:
                        media_formats.append(entry)

            if not media_formats:
                return jsonify({'error': 'No downloadable video found. Please check the URL and try again.'}), 404

            # Sort by quality (highest first), progressive first for same height
            sorted_media = sorted(media_formats, key=lambda x: (x.get('height', 0), 1 if x.get('video_only') else 2), reverse=True)

            # Generate filename based on platform
            uploader = info.get('uploader', info.get('uploader_id', 'unknown'))
            title = info.get('title', 'video')
            video_id = info.get('id', '')

            # Clean filename
            safe_title = re.sub(r'[^\w\-_\. ]', '_', title)[:50]
            safe_uploader = re.sub(r'[^\w\-_\. ]', '_', uploader)[:30]

            suggested_filename = f"{safe_uploader}_{safe_title}_{video_id}.mp4"

            return jsonify({
                'media': sorted_media,
                'bestAudio': best_audio,
                'suggestedFilename': suggested_filename,
                'title': title,
                'uploader': uploader,
                'duration': info.get('duration', 0)
            })

    except yt_dlp.utils.DownloadError as e:
        error_msg = str(e)
        if 'twitter' in url.lower() or 'x.com' in url.lower():
            return jsonify({'error': 'Invalid Twitter/X link. Please provide a proper link with a video.'}), 400
        elif 'youtube' in url.lower() or 'youtu.be' in url.lower():
            return jsonify({'error': 'Invalid YouTube link or video is unavailable.'}), 400
        else:
            return jsonify({'error': f'Unsupported platform or invalid link: {error_msg}'}), 400
    except Exception as e:
        return jsonify({'error': f'Failed to process URL: {str(e)}'}), 500

@csrf.exempt
@app.route('/proxy_merge_download')
def proxy_merge_download():
    video_url = request.args.get('video')
    audio_url = request.args.get('audio')
    filename = request.args.get('filename', f"merged_{int(time.time())}.mp4")
    if not video_url or not audio_url:
        return "Missing video or audio URL", 400

    # Download streams to temp files
    try:
        v_path = os.path.join(temp_dir, f"v_{uuid.uuid4()}.mp4")
        a_path = os.path.join(temp_dir, f"a_{uuid.uuid4()}.m4a")
        out_path = os.path.join(temp_dir, filename)

        vr = requests.get(video_url, stream=True)
        vr.raise_for_status()
        with open(v_path, 'wb') as vf:
            for chunk in vr.iter_content(chunk_size=8192):
                if chunk:
                    vf.write(chunk)

        ar = requests.get(audio_url, stream=True)
        ar.raise_for_status()
        with open(a_path, 'wb') as af:
            for chunk in ar.iter_content(chunk_size=8192):
                if chunk:
                    af.write(chunk)

        # Merge using ffmpeg
        subprocess.run(['ffmpeg', '-y', '-i', v_path, '-i', a_path, '-c', 'copy', out_path], check=True)
        return send_from_directory(temp_dir, os.path.basename(out_path), as_attachment=True)
    except Exception as e:
        return f"Failed to merge: {e}", 500

@csrf.exempt
@app.route('/proxy_download')
def proxy_download():
    url = request.args.get('url')
    filename = request.args.get('filename', 'download.mp4')
    if not url: return "Missing URL", 400
    try:
        r = requests.get(url, stream=True)
        r.raise_for_status()
        return Response(r.iter_content(chunk_size=8192),
                        mimetype=r.headers.get('Content-Type'),
                        headers={"Content-Disposition": f"attachment;filename=\"{filename}\""})
    except Exception as e:
        return f"Failed to fetch media: {e}", 500

@csrf.exempt
@app.route('/extract_highlights', methods=['POST'])
def extract_highlights_route():
    server_filename = request.form.get('serverFilename')
    if not server_filename: return jsonify({'error': 'No server file reference provided'}), 400
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], server_filename)
    # Prefer converted PDF when available
    pdf_path = get_pdf_for_serverfile(server_filename, filepath)
    if not pdf_path:
        return jsonify({'error': 'Could not obtain a PDF version of the uploaded file for highlight extraction.'}), 400
    try:
        highlights = extract_highlights(pdf_path)
        if not highlights:
            return jsonify({'error': 'No highlights were found in the PDF.'}), 400

        # Write notes PDF into temp upload folder
        pdf_filename = os.path.basename(pdf_path).replace('.pdf', '_notes.pdf')
        pdf_path_out = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
        create_modern_pdf(highlights, pdf_path_out)

        final_stats = get_doc_stats(pdf_path_out)

        # Create docx notes next to pdf
        docx_filename = os.path.basename(pdf_path).replace('.pdf', '_notes.docx')
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
        create_docx_from_highlights(highlights, docx_path)

        return jsonify({'previewUrl': f'/temp/{pdf_filename}', 'docxUrl': f'/temp/{docx_filename}', 'finalStats': final_stats})
    except Exception as e: return jsonify({'error': f'An unexpected error occurred: {str(e)}'}), 500

@csrf.exempt
@app.route('/add_header_footer', methods=['POST'])
def add_header_footer_route():
    form_data = request.form
    server_filename = form_data.get('serverFilename')
    if not server_filename: return jsonify({'error': 'No server file reference provided'}), 400
    input_filepath = os.path.join(app.config['UPLOAD_FOLDER'], server_filename)
    intermediate_pdf_path = get_pdf_for_serverfile(server_filename, input_filepath)
    if not intermediate_pdf_path:
        return jsonify({'error': 'Could not obtain a PDF version of the uploaded file. Try again later or upload a PDF directly.'}), 400
    
    headers = {'left': form_data.get('headerLeft', ''), 'center': form_data.get('headerCenter', ''), 'right': form_data.get('headerRight', '')}
    footers = {'left': form_data.get('footerLeft', ''), 'center': form_data.get('footerCenter', ''), 'right': form_data.get('footerRight', '')}
    page_num_placement = form_data.get('pageNumPlacement', 'footer-center')
    page_num_format = form_data.get('pageNumFormat', 'numeric')
    overlap_resolution = form_data.get('overlapResolution', 'after')
    margin_size = form_data.get('marginSize', 'normal')
    chapter_num = form_data.get('chapterNum', '1')
    page_num_enabled = form_data.get('isPageNumEnabled') == 'true'
    hf_enabled = form_data.get('isHfEnabled') == 'true'
    try: start_page_num = int(form_data.get('startPageNum', 1))
    except (ValueError, TypeError): start_page_num = 1
    
    output_filename = f"{uuid.uuid4()}_final.pdf"
    output_filepath = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
    
    try:
        add_header_footer_to_pdf(intermediate_pdf_path, output_filepath, headers, footers, start_page_num, page_num_placement, page_num_format, overlap_resolution, margin_size, chapter_num, page_num_enabled, hf_enabled)
        final_stats = get_doc_stats(output_filepath)
        # Also provide DOCX version
        docx_name = output_filename.replace('.pdf', '.docx')
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_name)
        create_docx_from_pdf(output_filepath, docx_path)
        return jsonify({'previewUrl': f'/temp/{output_filename}', 'docxUrl': f'/temp/{docx_name}', 'finalStats': final_stats})
    except Exception as e:
        return jsonify({'error': f'Failed to process PDF: {str(e)}'}), 500

port = int(os.environ.get("PORT", 5001))
app.run(host="0.0.0.0", port=port)

application = app
